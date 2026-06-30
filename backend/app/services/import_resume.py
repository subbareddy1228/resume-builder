import io
import json

from app.config import settings

EMPTY_CONTENT = {
    "contact": {"name": "", "email": "", "phone": "", "location": "", "linkedin": ""},
    "summary": "",
    "experience": [],
    "education": [],
    "skills": [],
    "projects": [],
    "certifications": [],
    "internships": [],
}


def extract_text_from_file(raw_bytes: bytes, ext: str) -> str:
    if ext == "pdf":
        return _extract_pdf_text(raw_bytes)
    elif ext == "docx":
        return _extract_docx_text(raw_bytes)
    raise ValueError(f"Unsupported extension: {ext}")


def _extract_pdf_text(raw_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw_bytes))
    text_parts = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def _extract_docx_text(raw_bytes: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull text out of tables, since some resumes use table layouts
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


PARSE_PROMPT = """You are an expert resume parser. Extract structured data from the resume text below
and return ONLY valid JSON matching this exact schema, nothing else — no markdown fences, no explanation.

Schema:
{{
  "contact": {{"name": "", "email": "", "phone": "", "location": "", "linkedin": ""}},
  "summary": "",
  "experience": [{{"id": "", "company": "", "role": "", "start": "", "end": "", "bullets": [""]}}],
  "education": [{{"id": "", "school": "", "degree": "", "year": ""}}],
  "skills": [""],
  "projects": [{{"id": "", "name": "", "description": "", "tech": "", "link": ""}}],
  "certifications": [{{"id": "", "name": "", "issuer": "", "year": ""}}],
  "internships": [{{"id": "", "company": "", "role": "", "start": "", "end": "", "bullets": [""]}}]
}}

Rules:
- Generate a short random alphanumeric string for every "id" field (e.g. "a1b2c3").
- If a field isn't present in the resume, use an empty string "" or empty array [], never null.
- Distinguish "internships" from "experience" only if explicitly labeled as an internship; otherwise put all work history under "experience".
- Keep bullet points close to the original wording — do not invent achievements.
- "skills" should be a flat list of individual skill strings, splitting comma/bullet separated skill lists.
- Dates: use whatever format appears in the resume (e.g. "Jan 2022", "2022").

Resume text:
<resume_text>
{resume_text}
</resume_text>

Return ONLY the JSON object."""


def parse_resume_with_claude(resume_text: str) -> dict:
    import anthropic
    client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    # Cap input length defensively
    trimmed = resume_text[:15000]

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": PARSE_PROMPT.format(resume_text=trimmed)}],
    )

    raw = "".join(block.text for block in message.content if block.type == "text").strip()

    # Strip accidental markdown fences just in case
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    parsed = json.loads(raw)

    # Merge onto EMPTY_CONTENT to guarantee every key exists
    result = dict(EMPTY_CONTENT)
    result.update({k: v for k, v in parsed.items() if k in EMPTY_CONTENT})
    return result
